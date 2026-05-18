import os
import io
import re
import tempfile
from pathlib import Path
from flask import Flask, request, jsonify, send_file, render_template

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 50 MB

SUPPORTED_TYPES = {
    ".pdf": "PDF",
    ".docx": "Word Document",
    ".doc": "Word Document (legacy)",
    ".txt": "Plain Text",
    ".html": "HTML",
    ".htm": "HTML",
    ".rtf": "Rich Text Format",
    ".md": "Markdown",
    ".csv": "CSV",
}


def convert_txt(content_bytes: bytes) -> str:
    text = content_bytes.decode("utf-8", errors="replace")
    lines = text.splitlines()
    output = []
    for line in lines:
        stripped = line.rstrip()
        output.append(stripped)
    return "\n".join(output)


def convert_md(content_bytes: bytes) -> str:
    return content_bytes.decode("utf-8", errors="replace")


def convert_html(content_bytes: bytes) -> str:
    from markdownify import markdownify as md
    html = content_bytes.decode("utf-8", errors="replace")
    result = md(html, heading_style="ATX", bullets="-", strip=["script", "style"])
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip()


def convert_docx(content_bytes: bytes) -> str:
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(content_bytes))
    output = []

    HEADING_MAP = {
        "Heading 1": "#",
        "Heading 2": "##",
        "Heading 3": "###",
        "Heading 4": "####",
        "Heading 5": "#####",
        "Heading 6": "######",
    }

    def para_to_md(para):
        style_name = para.style.name if para.style else ""
        prefix = HEADING_MAP.get(style_name, "")

        runs_md = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            if run.underline:
                text = f"<u>{text}</u>"
            runs_md.append(text)

        line = "".join(runs_md)

        # List detection
        numPr = para._element.find(qn("w:pPr"))
        is_list = False
        list_level = 0
        if numPr is not None:
            numPrEl = numPr.find(qn("w:numPr"))
            if numPrEl is not None:
                is_list = True
                ilvl = numPrEl.find(qn("w:ilvl"))
                if ilvl is not None:
                    list_level = int(ilvl.get(qn("w:val"), "0"))

        if prefix:
            return f"{prefix} {line}" if line else ""
        elif is_list:
            indent = "  " * list_level
            return f"{indent}- {line}"
        else:
            return line

    for para in doc.paragraphs:
        md_line = para_to_md(para)
        if md_line is not None:
            output.append(md_line)

    # Tables
    for table in doc.tables:
        if not table.rows:
            continue
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)

        if rows:
            header = "| " + " | ".join(rows[0]) + " |"
            sep = "| " + " | ".join(["---"] * len(rows[0])) + " |"
            body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
            output.append("")
            output.append(header)
            output.append(sep)
            if body:
                output.append(body)
            output.append("")

    result = "\n".join(output)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip()


def convert_pdf(content_bytes: bytes) -> str:
    from pdfminer.high_level import extract_pages
    from pdfminer.layout import LTTextContainer, LTChar, LTAnon, LTFigure

    output = []
    prev_font_size = None

    for page_layout in extract_pages(io.BytesIO(content_bytes)):
        for element in page_layout:
            if isinstance(element, LTTextContainer):
                text = element.get_text().strip()
                if not text:
                    continue

                # Detect approximate font size for heading heuristic
                font_sizes = []
                for text_line in element:
                    for char in text_line:
                        if isinstance(char, LTChar):
                            font_sizes.append(char.size)

                avg_size = sum(font_sizes) / len(font_sizes) if font_sizes else 12

                if avg_size >= 20:
                    output.append(f"\n# {text}\n")
                elif avg_size >= 16:
                    output.append(f"\n## {text}\n")
                elif avg_size >= 14:
                    output.append(f"\n### {text}\n")
                else:
                    output.append(text)

    result = "\n".join(output)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip()


def convert_rtf(content_bytes: bytes) -> str:
    from striprtf.striprtf import rtf_to_text
    text = rtf_to_text(content_bytes.decode("utf-8", errors="replace"))
    lines = text.splitlines()
    output = [line.rstrip() for line in lines]
    return "\n".join(output).strip()


def convert_csv(content_bytes: bytes) -> str:
    import csv
    text = content_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return ""

    col_count = max(len(r) for r in rows)
    for r in rows:
        while len(r) < col_count:
            r.append("")

    header = "| " + " | ".join(rows[0]) + " |"
    sep = "| " + " | ".join(["---"] * col_count) + " |"
    body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
    return f"{header}\n{sep}\n{body}"


CONVERTERS = {
    ".txt": convert_txt,
    ".md": convert_md,
    ".html": convert_html,
    ".htm": convert_html,
    ".docx": convert_docx,
    ".pdf": convert_pdf,
    ".rtf": convert_rtf,
    ".csv": convert_csv,
}


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/convert", methods=["POST"])
def convert():
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "No file selected"}), 400

    ext = Path(file.filename).suffix.lower()

    if ext not in CONVERTERS:
        supported = ", ".join(sorted(CONVERTERS.keys()))
        return jsonify({
            "error": f"Unsupported file type '{ext}'. Supported: {supported}"
        }), 400

    content_bytes = file.read()
    original_name = Path(file.filename).stem

    try:
        converter = CONVERTERS[ext]
        md_content = converter(content_bytes)
    except Exception as e:
        return jsonify({"error": f"Conversion failed: {str(e)}"}), 500

    if request.args.get("preview") == "1":
        return jsonify({"content": md_content, "filename": f"{original_name}.md"})

    md_bytes = md_content.encode("utf-8")
    buf = io.BytesIO(md_bytes)
    buf.seek(0)

    return send_file(
        buf,
        mimetype="text/markdown",
        as_attachment=True,
        download_name=f"{original_name}.md",
    )


if __name__ == "__main__":
    app.run(debug=True, port=5050)
